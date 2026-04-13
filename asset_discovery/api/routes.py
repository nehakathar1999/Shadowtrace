from fastapi import APIRouter, Query, HTTPException, Body, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel
from scanners.ping_sweep import threaded_ping_sweep
from scanners.port_scanner import threaded_port_scan
from scanners.service_detector import detect_services
from intelligence.device_classifier import classify_device
from intelligence.os_fingerprint import detect_os_details
from intelligence.protocol_detector import detect_insecure_protocols
from intelligence.tls_analyzer import analyze_tls
from intelligence.vuln_assessor import assess_vulnerabilities
from intelligence.owasp_top10 import scan_owasp_top_10
from intelligence.product_analysis import executive_summary
from intelligence.web_advanced import (
    run_analyze_api_document,
    run_assess_api_endpoints,
    run_discover_web_surface,
)
from scheduler.scan_scheduler import (
    compare_reports,
    delete_schedule,
    list_schedule_runs,
    list_schedules,
    save_schedule,
    schedule_scan_job,
    start_scheduler,
)
from utils.network_info import get_hostname, get_mac, get_vendor
import ipaddress
import socket
import os
import json
import re
import uuid
import html
import hashlib
import hmac
import time
import threading
import asyncio
import secrets
from datetime import datetime, timedelta, timezone
from config.settings import settings
from database.db import SessionLocal
from database.models import (
    Asset as AssetModel,
    InsecureProtocol,
    LoginEvent,
    OWASPFinding,
    OWASPResult,
    Port as PortModel,
    Scan as ScanModel,
    SignupEvent,
    TLSIssue,
    User as UserModel,
    Vulnerability as VulnerabilityModel,
)

try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image
    from reportlab.pdfgen import canvas
except ImportError:
    A4 = None
    letter = None
    colors = None
    getSampleStyleSheet = None
    ParagraphStyle = None
    TA_LEFT = None
    TA_CENTER = None
    inch = None
    SimpleDocTemplate = None
    Table = None
    TableStyle = None
    Paragraph = None
    Spacer = None
    PageBreak = None
    Image = None
    canvas = None

try:
    from reportlab.graphics.shapes import Drawing, Circle
    from reportlab.graphics.charts.piecharts import Pie
except ImportError:
    Drawing = None
    Circle = None
    Pie = None

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

router = APIRouter()
SCAN_JOBS = {}
SCAN_JOBS_LOCK = threading.Lock()
PASSWORD_HASH_ITERATIONS = 100_000


class SignupRequest(BaseModel):
    name: str
    email: str
    password: str


class LoginRequest(BaseModel):
    email: str
    password: str


class ForgotPasswordRequest(BaseModel):
    email: str
    new_password: str


def _normalize_email(email: str) -> str:
    return str(email or "").strip().lower()


def _is_valid_email(email: str) -> bool:
    return bool(re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email))


def _hash_password(password: str, salt: str | None = None) -> str:
    salt_value = salt or secrets.token_hex(16)
    derived = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt_value.encode("utf-8"),
        PASSWORD_HASH_ITERATIONS,
    )
    return f"{salt_value}${derived.hex()}"


def _verify_password(password: str, stored_hash: str) -> bool:
    try:
        salt, existing_hash = stored_hash.split("$", 1)
    except ValueError:
        return False
    candidate_hash = _hash_password(password, salt).split("$", 1)[1]
    return hmac.compare_digest(existing_hash, candidate_hash)


def _serialize_user(user: UserModel) -> dict:
    return {
        "id": user.id,
        "name": user.name,
        "email": user.email,
        "created_at": user.created_at.isoformat() + "Z" if user.created_at else None,
    }


def _client_ip(request: Request) -> str | None:
    forwarded_for = request.headers.get("x-forwarded-for", "")
    if forwarded_for:
        return forwarded_for.split(",")[0].strip() or None
    if request.client:
        return request.client.host
    return None


def _user_agent(request: Request) -> str | None:
    return request.headers.get("user-agent") or None


def _job_log(job_id: str, message: str):
    with SCAN_JOBS_LOCK:
        job = SCAN_JOBS.get(job_id)
        if not job:
            return
        job["logs"].append({
            "time": datetime.now().strftime("[%I:%M:%S %p]").lower(),
            "msg": message,
        })
        job["updated_at"] = datetime.utcnow().isoformat() + "Z"


def _job_update(job_id: str, **fields):
    with SCAN_JOBS_LOCK:
        job = SCAN_JOBS.get(job_id)
        if not job:
            return
        job.update(fields)
        job["updated_at"] = datetime.utcnow().isoformat() + "Z"


def _job_wait_if_paused(job_id: str):
    while True:
        with SCAN_JOBS_LOCK:
            job = SCAN_JOBS.get(job_id)
            if not job:
                return False
            if job.get("cancel_requested"):
                job["status"] = "cancelled"
                job["updated_at"] = datetime.utcnow().isoformat() + "Z"
                return False
            if not job.get("pause_requested"):
                if job.get("status") == "paused":
                    job["status"] = "running"
                job["updated_at"] = datetime.utcnow().isoformat() + "Z"
                return True
            job["status"] = "paused"
            job["updated_at"] = datetime.utcnow().isoformat() + "Z"
        time.sleep(0.25)


def _job_snapshot(job: dict) -> dict:
    return {
        "job_id": job["job_id"],
        "target": job["target"],
        "status": job["status"],
        "progress": job["progress"],
        "stage_index": job["stage_index"],
        "stage_label": job["stage_label"],
        "logs": job["logs"],
        "result": job.get("result"),
        "error": job.get("error"),
        "created_at": job["created_at"],
        "updated_at": job["updated_at"],
    }


@router.post("/auth/signup")
def signup(payload: SignupRequest, request: Request):
    name = str(payload.name or "").strip()
    email = _normalize_email(payload.email)
    password = str(payload.password or "")

    if len(name) < 2:
        raise HTTPException(status_code=400, detail="Name must be at least 2 characters")
    if not _is_valid_email(email):
        raise HTTPException(status_code=400, detail="Please enter a valid email")
    if len(password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")

    db = SessionLocal()
    try:
        existing = db.query(UserModel).filter(UserModel.email == email).first()
        if existing:
            raise HTTPException(status_code=409, detail="An account with this email already exists")

        user = UserModel(
            name=name,
            email=email,
            password_hash=_hash_password(password),
        )
        db.add(user)
        db.flush()

        db.add(
            SignupEvent(
                user_id=user.id,
                name=user.name,
                email=user.email,
                ip_address=_client_ip(request),
                user_agent=_user_agent(request),
            )
        )
        db.commit()
        db.refresh(user)
        return {"success": True, "user": _serialize_user(user)}
    finally:
        db.close()


@router.post("/auth/login")
def login(payload: LoginRequest, request: Request):
    email = _normalize_email(payload.email)
    password = str(payload.password or "")

    if not _is_valid_email(email):
        raise HTTPException(status_code=400, detail="Please enter a valid email")
    if not password:
        raise HTTPException(status_code=400, detail="Password is required")

    db = SessionLocal()
    try:
        user = db.query(UserModel).filter(UserModel.email == email).first()
        if not user or not _verify_password(password, user.password_hash):
            raise HTTPException(status_code=401, detail="Invalid email or password")

        db.add(
            LoginEvent(
                user_id=user.id,
                email=user.email,
                ip_address=_client_ip(request),
                user_agent=_user_agent(request),
            )
        )
        db.commit()
        return {"success": True, "user": _serialize_user(user)}
    finally:
        db.close()


@router.post("/auth/forgot-password")
def forgot_password(payload: ForgotPasswordRequest):
    email = _normalize_email(payload.email)
    new_password = str(payload.new_password or "")

    if not _is_valid_email(email):
        raise HTTPException(status_code=400, detail="Please enter a valid email")
    if len(new_password) < 6:
        raise HTTPException(status_code=400, detail="New password must be at least 6 characters")

    db = SessionLocal()
    try:
        user = db.query(UserModel).filter(UserModel.email == email).first()
        if not user:
            raise HTTPException(status_code=404, detail="No account found with this email")
        user.password_hash = _hash_password(new_password)
        db.commit()
        return {"success": True, "message": "Password updated successfully"}
    finally:
        db.close()


def resolve_domain(domain: str):
    """Resolve a domain to one or more IPv4 addresses."""

    try:
        infos = socket.getaddrinfo(domain, None, family=socket.AF_INET)
        ips = [info[4][0] for info in infos]
        return list(dict.fromkeys(ips))
    except Exception:
        return []


def is_domain_input(target: str) -> bool:
    """Return True if the given target string appears to be a domain name."""

    # Exclude explicit CIDR/ranges/multiple targets
    if "," in target or "-" in target:
        return False

    try:
        ipaddress.ip_address(target)
        return False
    except ValueError:
        pass

    try:
        ipaddress.ip_network(target, strict=False)
        return False
    except ValueError:
        pass

    return True


def _sanitize_report_path(path: str) -> str:
    """Sanitize report path to prevent directory traversal."""
    if not path or ".." in path or path.startswith("/") or path.startswith("\\"):
        raise HTTPException(status_code=400, detail="Invalid report path")

    normalized = os.path.normpath(path)
    if not normalized.startswith("reports"):
        raise HTTPException(status_code=400, detail="Invalid report path")

    return normalized


def _format_report_filename(target: str, format: str) -> str:
    base_name = os.path.splitext(os.path.basename(target))[0]
    return f"{base_name}.{format}"


def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


OWASP_TOP_10_2025 = [
    ("A01:2025", "Broken Access Control"),
    ("A02:2025", "Security Misconfiguration"),
    ("A03:2025", "Software Supply Chain Failures"),
    ("A04:2025", "Cryptographic Failures"),
    ("A05:2025", "Injection"),
    ("A06:2025", "Insecure Design"),
    ("A07:2025", "Authentication Failures"),
    ("A08:2025", "Software or Data Integrity Failures"),
    ("A09:2025", "Security Logging and Alerting Failures"),
    ("A10:2025", "Mishandling of Exceptional Conditions"),
]


def _repo_root() -> str:
    return os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))


IST = timezone(timedelta(hours=5, minutes=30))


def _find_brand_logo() -> str | None:
    root = _repo_root()
    candidates = [
        os.path.join(root, "frontend", "src", "assets", "shadowtrace.png"),
        # os.path.join(root, "frontend", "src", "assets", "LOGO_CYBERSPACE.png"),
        # os.path.join(root, "frontend", "src", "assets", "kristellar-cyberspace.png"),
        # os.path.join(root, "frontend", "src", "assets", "kristellar-cyberspace.jpg"),
        # os.path.join(root, "frontend", "src", "assets", "kristellar-logo-white.png"),
        # os.path.join(root, "frontend", "src", "assets", "kristellar-logo-white.jpg"),
        # os.path.join(root, "frontend", "src", "assets", "kristellar-aerospace.png"),
        # os.path.join(root, "frontend", "src", "assets", "kristellar-aerospace.jpg"),
        # os.path.join(root, "frontend", "src", "assets", "kristellar.png"),
        # os.path.join(root, "frontend", "src", "assets", "kristellar.jpg"),
        # os.path.join(root, "frontend", "public", "logo.png"),
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def _safe_text(value, default="N/A") -> str:
    text = str(value).strip() if value is not None else ""
    return text or default


def _safe_paragraph(text: str, style):
    return Paragraph(html.escape(_safe_text(text)).replace("\n", "<br/>"), style)


def _severity_rank(severity: str) -> int:
    order = {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1, "INFO": 0}
    return order.get(str(severity or "").upper(), -1)


def _severity_color(severity: str):
    mapping = {
        "CRITICAL": colors.HexColor("#c62828"),
        "HIGH": colors.HexColor("#ef6c00"),
        "MEDIUM": colors.HexColor("#f9a825"),
        "LOW": colors.HexColor("#43a047"),
        "INFO": colors.HexColor("#1e88e5"),
    }
    return mapping.get(str(severity or "").upper(), colors.HexColor("#616161"))


def _normalize_status(vulnerability: dict) -> str:
    status = str(vulnerability.get("status") or "").strip()
    return status if status else "Open"


def _collect_vulnerability_rows(payload: dict) -> list[dict]:
    rows = []
    for asset in payload.get("assets", []):
        host = asset.get("domain") or asset.get("hostname") or asset.get("resolved_ip") or asset.get("ip") or "Unknown"
        ip = asset.get("resolved_ip") or asset.get("ip") or "Unknown"
        for vulnerability in asset.get("vulnerabilities", []):
            rows.append({
                "host": host,
                "ip": ip,
                "port": vulnerability.get("port", "-"),
                "service": vulnerability.get("service") or "unknown",
                "product": vulnerability.get("product") or "Unknown product",
                "version": vulnerability.get("version") or "-",
                "cve": vulnerability.get("cve") or "N/A",
                "severity": (vulnerability.get("severity") or "UNKNOWN").upper(),
                "cvss_score": vulnerability.get("cvss_score"),
                "title": vulnerability.get("title") or "Unnamed vulnerability",
                "description": vulnerability.get("description") or "No description available.",
                "remediation": vulnerability.get("remediation") or "Apply the vendor fix, reduce exposure, and verify the remediation with a rescan.",
                "status": _normalize_status(vulnerability),
                "confidence": vulnerability.get("confidence") or "unknown",
                "confidence_score": vulnerability.get("confidence_score"),
                "validation_state": vulnerability.get("validation_state") or "unknown",
                "proof": vulnerability.get("proof") or {},
                "business_impact": vulnerability.get("business_impact") or "No business impact summary available.",
                "compliance_mapping": vulnerability.get("compliance_mapping") or [],
            })
    rows.sort(
        key=lambda row: (
            _severity_rank(row["severity"]),
            float(row["cvss_score"]) if isinstance(row["cvss_score"], (int, float)) else -1,
        ),
        reverse=True,
    )
    return rows


def _collect_port_rows(payload: dict) -> list[list[str]]:
    rows = []
    for asset in payload.get("assets", []):
        host = asset.get("domain") or asset.get("hostname") or asset.get("ip") or "Unknown"
        for port_info in asset.get("open_ports", []):
            rows.append([
                host,
                str(port_info.get("port", "-")),
                _safe_text(port_info.get("protocol"), "-").upper(),
                _safe_text(port_info.get("service"), "-"),
                _safe_text(port_info.get("product"), "-"),
                _safe_text(port_info.get("version"), "-"),
            ])
    return rows or [["No open port data available", "-", "-", "-", "-", "-"]]


def _collect_os_rows(payload: dict) -> list[list[str]]:
    rows = []
    for asset in payload.get("assets", []):
        host = asset.get("domain") or asset.get("hostname") or asset.get("ip") or "Unknown"
        rows.append([
            host,
            _safe_text(asset.get("os_name") or asset.get("os")),
            _safe_text(asset.get("os_family")),
            _safe_text(asset.get("os_accuracy")),
            _safe_text(asset.get("os_source")),
        ])
    return rows or [["No OS detection data available", "-", "-", "-", "-"]]


def _collect_protocol_rows(payload: dict, field_name: str) -> list[list[str]]:
    rows = []
    for asset in payload.get("assets", []):
        host = asset.get("domain") or asset.get("hostname") or asset.get("ip") or "Unknown"
        for item in asset.get(field_name, []):
            rows.append([
                host,
                str(item.get("port", "-")),
                _safe_text(item.get("protocol") or item.get("tls_version"), "-"),
                _safe_text(item.get("msg") or item.get("message"), "-"),
            ])
    label = "observations" if field_name == "tls_issues" else "findings"
    return rows or [[f"No {label} detected", "-", "-", "-"]]


def _collect_owasp_results(payload: dict) -> list[dict]:
    owasp = payload.get("owasp_top_10") or {}
    results = owasp.get("results")
    return results if isinstance(results, list) else []


def _collect_owasp_summary(payload: dict) -> dict:
    owasp = payload.get("owasp_top_10") or {}
    summary = owasp.get("summary")
    if isinstance(summary, dict):
        return summary
    return {
        "total_categories": 0,
        "categories_with_findings": 0,
        "total_findings": 0,
    }


def _owasp_findings_count(payload: dict) -> int:
    return sum(len(item.get("findings") or []) for item in _collect_owasp_results(payload))


def _should_include_owasp_section(payload: dict) -> bool:
    owasp = payload.get("owasp_top_10") or {}
    return bool(owasp.get("enabled"))


def _run_owasp_scan(target: str, include_domain: bool, auth_context: dict | None = None) -> dict:
    if not include_domain:
        return {
            "enabled": False,
            "target": target,
            "normalized_url": None,
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "summary": {
                "total_categories": len(OWASP_TOP_10_2025),
                "categories_with_findings": 0,
                "total_findings": 0,
            },
            "results": [],
            "note": "OWASP web checks run only for domain or URL based targets.",
        }

    try:
        return asyncio.run(scan_owasp_top_10(target, auth_context))
    except Exception as exc:
        return {
            "enabled": True,
            "target": target,
            "normalized_url": target,
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "summary": {
                "total_categories": len(OWASP_TOP_10_2025),
                "categories_with_findings": 0,
                "total_findings": 0,
            },
            "results": [],
            "error": str(exc),
        }


def _severity_totals(vulnerabilities: list[dict]) -> dict:
    totals = {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0}
    for vulnerability in vulnerabilities:
        severity = vulnerability["severity"]
        if severity in totals:
            totals[severity] += 1
    return totals


def _impact_text(vulnerability: dict) -> str:
    if vulnerability.get("business_impact"):
        return vulnerability["business_impact"]
    severity = vulnerability["severity"]
    title = vulnerability["title"]
    if severity == "CRITICAL":
        return f"{title} may enable full service compromise, remote code execution, or unauthorized high-impact access if exploited."
    if severity == "HIGH":
        return f"{title} may expose sensitive data, weaken perimeter controls, or enable meaningful lateral movement."
    if severity == "MEDIUM":
        return f"{title} increases the attack surface and may support chained exploitation when combined with other weaknesses."
    return f"{title} should still be remediated because it contributes to unnecessary exposure and defense-in-depth gaps."


def _suggested_fixes_text(vulnerability: dict) -> str:
    remediation = vulnerability["remediation"]
    return remediation + " Validate the affected service version, restrict external exposure where possible, and confirm closure with a repeat scan."


def _proof_summary(vulnerability: dict) -> str:
    proof = vulnerability.get("proof") or {}
    observed = proof.get("observed") or []
    request = proof.get("request") or {}
    response = proof.get("response") or {}
    payload = proof.get("payload")
    conclusion = proof.get("conclusion")

    parts = []
    if request:
        method = request.get("method") or request.get("protocol") or "request"
        target = request.get("url") or request.get("target_port") or vulnerability.get("port")
        parts.append(f"Request: {method} {target}")
    if payload:
        parts.append(f"Payload: {payload}")
    if response:
        status = response.get("status_code")
        service = response.get("service")
        version = response.get("version")
        detail = status if status is not None else f"{service or ''} {version or ''}".strip()
        if detail:
            parts.append(f"Response: {detail}")
    if observed:
        parts.append(f"Evidence: {'; '.join(str(item) for item in observed[:3])}")
    if conclusion:
        parts.append(f"Conclusion: {conclusion}")
    return " ".join(parts) if parts else "No proof artifacts were recorded for this finding."


def _additional_references(vulnerability: dict) -> list[str]:
    references = ["https://owasp.org/Top10/2025/"]
    cve_id = vulnerability.get("cve")
    if cve_id and cve_id != "N/A":
        references.append(f"https://nvd.nist.gov/vuln/detail/{cve_id}")
        references.append(f"https://www.cve.org/CVERecord?id={cve_id}")
    return references


def _build_report_styles():
    base = getSampleStyleSheet()
    return {
        "cover_date": ParagraphStyle("CoverDate", parent=base["Normal"], fontName="Helvetica", fontSize=12, textColor=colors.white, leading=16),
        "cover_kicker": ParagraphStyle("CoverKicker", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=12, leading=16, textColor=colors.HexColor("#e3b93f"), spaceAfter=8),
        "cover_title": ParagraphStyle("CoverTitle", parent=base["Title"], fontName="Helvetica-Bold", fontSize=44, leading=46, textColor=colors.white, spaceAfter=8),
        "cover_subtitle": ParagraphStyle("CoverSubtitle", parent=base["Normal"], fontName="Helvetica", fontSize=14, leading=18, textColor=colors.HexColor("#b6d6ff")),
        "cover_accent": ParagraphStyle("CoverAccent", parent=base["Title"], fontName="Helvetica-Bold", fontSize=22, leading=28, textColor=colors.HexColor("#4aa8e8")),
        "brand": ParagraphStyle("Brand", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=18, textColor=colors.HexColor("#cbd5e1"), alignment=TA_LEFT),
        "header_bar": ParagraphStyle("HeaderBar", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=colors.white),
        "section_title": ParagraphStyle("SectionTitle", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=colors.HexColor("#1f3f77"), spaceAfter=8),
        "subsection": ParagraphStyle("Subsection", parent=base["Heading3"], fontName="Helvetica-Bold", fontSize=14, leading=18, textColor=colors.HexColor("#f59e0b"), spaceAfter=6),
        "body": ParagraphStyle("Body", parent=base["BodyText"], fontName="Helvetica", fontSize=10, leading=14, textColor=colors.HexColor("#111827")),
        "body_small": ParagraphStyle("BodySmall", parent=base["BodyText"], fontName="Helvetica", fontSize=9, leading=12, textColor=colors.HexColor("#111827")),
        "metric_label": ParagraphStyle("MetricLabel", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=12, textColor=colors.HexColor("#6b7280"), alignment=TA_CENTER),
        "metric_value": ParagraphStyle("MetricValue", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=17, leading=20, textColor=colors.HexColor("#111827"), alignment=TA_CENTER),
        "toc_item": ParagraphStyle("TocItem", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=13, leading=18, textColor=colors.HexColor("#111827")),
        "closing_contact": ParagraphStyle("ClosingContact", parent=base["Normal"], fontName="Helvetica", fontSize=11, leading=15, textColor=colors.white),
        "closing_link": ParagraphStyle("ClosingLink", parent=base["Normal"], fontName="Helvetica", fontSize=11, leading=15, textColor=colors.HexColor("#111827")),
        "cover_logo_word": ParagraphStyle("CoverLogoWord", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=24, leading=26, textColor=colors.white),
        "cover_logo_tag": ParagraphStyle("CoverLogoTag", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=11, leading=13, textColor=colors.white),
        "closing_logo_word": ParagraphStyle("ClosingLogoWord", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=16, leading=18, textColor=colors.white),
        "closing_logo_tag": ParagraphStyle("ClosingLogoTag", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=colors.white),
    }


def _section_banner(text: str, styles) -> Table:
    table = Table([[Paragraph(html.escape(text), styles["header_bar"])]], colWidths=[510])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#1f3f77")),
        ("LEFTPADDING", (0, 0), (-1, -1), 18),
        ("RIGHTPADDING", (0, 0), (-1, -1), 18),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
    ]))
    return table


def _styled_table(rows: list[list], col_widths: list[int], header_background=None) -> Table:
    header_background = header_background or colors.HexColor("#1f3f77")
    table = Table(rows, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), header_background),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("LEADING", (0, 0), (-1, 0), 11),
        ("BACKGROUND", (0, 1), (-1, -1), colors.white),
        ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor("#111827")),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 1), (-1, -1), 8.5),
        ("LEADING", (0, 1), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


def _donut_graph_summary(segments: list[tuple[str, int, str]], styles, width: float):
    if Drawing is None or Pie is None or Circle is None:
        rows = [["Category", "Findings"]]
        if segments:
            for label, count, _ in segments:
                shown_count = 0 if label == "No Findings" else count
                rows.append([label, str(shown_count)])
        else:
            rows.append(["No Findings", "0"])
        return _styled_table(rows, [250, 120])
    if not segments:
        segments = [("No Findings", 1, "#94a3b8")]

    drawing = Drawing(132, 122)
    pie = Pie()
    pie.x = 10
    pie.y = 10
    pie.width = 104
    pie.height = 104
    pie.data = [segment[1] for segment in segments]
    pie.labels = [""] * len(segments)
    pie.sideLabels = 0
    pie.startAngle = 110
    pie.slices.strokeColor = colors.white
    pie.slices.strokeWidth = 0.5
    for i, segment in enumerate(segments):
        pie.slices[i].fillColor = colors.HexColor(segment[2])
    drawing.add(pie)
    # Create a donut look using a white center circle.
    drawing.add(Circle(62, 60, 28, fillColor=colors.white, strokeColor=colors.white))

    legend_rows = []
    for label, count, hex_color in segments:
        shown_count = 0 if label == "No Findings" else count
        legend_rows.append([
            Paragraph(f'<font color="{hex_color}">&#9632;</font>', styles["body_small"]),
            Paragraph(f"{label}: {shown_count}", styles["body_small"]),
        ])
    legend = Table(legend_rows, colWidths=[12, max(width - 144, 92)])
    legend.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    chart_layout = Table([[drawing, legend]], colWidths=[136, width - 136])
    chart_layout.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    return chart_layout


def _severity_graph_summary(severity_totals: dict, styles, width: float):
    severity_defs = [
        ("CRITICAL", "Critical", "#c62828"),
        ("HIGH", "High", "#ef6c00"),
        ("MEDIUM", "Medium", "#f9a825"),
        ("LOW", "Low", "#43a047"),
    ]
    segments = []
    for key, label, hex_color in severity_defs:
        count = int(severity_totals.get(key, 0) or 0)
        if count > 0:
            segments.append((label, count, hex_color))

    if Drawing is None or Pie is None or Circle is None:
        severity_rows = [["Severity", "Findings"]]
        for severity in ("CRITICAL", "HIGH", "MEDIUM", "LOW"):
            severity_rows.append([severity.title(), str(severity_totals[severity])])
        return _styled_table(severity_rows, [250, 120])

    return _donut_graph_summary(segments, styles, width)


def _owasp_graph_summary(owasp_results: list[dict], styles, width: float):
    owasp_palette = [
        "#7c3aed",
        "#2563eb",
        "#0891b2",
        "#0f766e",
        "#65a30d",
        "#ca8a04",
        "#ea580c",
        "#dc2626",
        "#db2777",
        "#4f46e5",
    ]
    segments = []
    color_index = 0
    for item in owasp_results:
        count = int(item.get("findings_count", 0) or 0)
        if count <= 0:
            continue
        label = item.get("title") or item.get("id") or "OWASP Finding"
        segments.append((label, count, owasp_palette[color_index % len(owasp_palette)]))
        color_index += 1

    if Drawing is None or Pie is None or Circle is None:
        owasp_rows = [["OWASP Category", "Findings"]]
        if segments:
            for label, count, _ in segments:
                owasp_rows.append([label, str(count)])
        else:
            owasp_rows.append(["No Findings", "0"])
        return _styled_table(owasp_rows, [250, 120])

    return _donut_graph_summary(segments, styles, width)


def _draw_page_footer(canv, doc):
    page_width = doc.pagesize[0]
    page_number = canv.getPageNumber()
    if page_number <= 1:
        return
    canv.saveState()
    canv.setStrokeColor(colors.HexColor("#98b5de"))
    canv.setLineWidth(0.8)
    canv.line(28, 34, page_width - 28, 34)
    canv.setFont("Helvetica", 8.5)
    canv.setFillColor(colors.HexColor("#1f3f77"))
    canv.drawString(28, 20, "Kristellar Aerospace Pvt Ltd | Confidential")
    canv.drawRightString(page_width - 28, 20, f"{page_number:02d}")
    canv.restoreState()


def _draw_body_page(canv, doc):
    canv.saveState()
    page_width, page_height = doc.pagesize
    canv.setFillColor(colors.HexColor("#1f3f77"))
    canv.rect(0, page_height - 22, page_width, 22, stroke=0, fill=1)
    canv.restoreState()
    _draw_page_footer(canv, doc)


def _generate_pdf_report(payload: dict, output_path: str):
    if SimpleDocTemplate is not None and A4 is not None:
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=42,
            leftMargin=42,
            topMargin=34,
            bottomMargin=34,
        )
        styles = _build_report_styles()
        summary = payload.get("vulnerability_summary", {})
        product_summary = payload.get("product_summary", {})
        owasp_summary = _collect_owasp_summary(payload)
        include_owasp_section = _should_include_owasp_section(payload)
        generated_at_ist = datetime.now(IST)
        generated_at = generated_at_ist.strftime("%Y-%m-%d %I:%M:%S %p IST")
        vulnerabilities = _collect_vulnerability_rows(payload)
        owasp_results = _collect_owasp_results(payload)
        severity_totals = _severity_totals(vulnerabilities)
        logo_path = _find_brand_logo()
        story = []

        cover_content = [
            Spacer(1, 8),
        ]
        if logo_path:
            cover_content.append(Image(logo_path, width=2.1 * inch, height=0.9 * inch, kind="proportional"))
            cover_content.append(Spacer(1, 18))
        else:
            cover_content.append(Spacer(1, 6))
        cover_content.extend([
            Paragraph(generated_at_ist.strftime("%B %d, %Y"), styles["cover_date"]),
            Spacer(1, 42),
            Paragraph("SECURITY AUDIT", styles["cover_title"]),
            Paragraph("REPORT", styles["cover_title"]),
            Spacer(1, 14),
            Paragraph(f"{html.escape(_safe_text(payload.get('input')).upper())}", styles["cover_accent"]),
            Spacer(1, 24),
        ])
        cover_content.append(Spacer(1, 300))
        cover_content.extend([
            Paragraph("Kristellar Aerospace", styles["brand"]),
            Spacer(1, 8),
            _safe_paragraph(f"Generated on {generated_at}", styles["cover_date"]),
        ])
        cover = Table([[cover_content]], colWidths=[doc.width], rowHeights=[doc.height - 44])
        cover.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#0f2244")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 24),
            ("RIGHTPADDING", (0, 0), (-1, -1), 24),
            ("TOPPADDING", (0, 0), (-1, -1), 24),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 24),
        ]))
        story.extend([cover, PageBreak()])

        toc_rows = [
            [Paragraph("<b>1.</b> Executive Summary", styles["toc_item"])],
            [Paragraph("1.1 Scope of Testing", styles["body"])],
            [Paragraph("1.2 Graphical Summary", styles["body"])],
            [Paragraph("1.3 List of Vulnerabilities", styles["body"])],
            [Paragraph("<b>2.</b> Technical Findings", styles["toc_item"])],
            [Paragraph("2.1 Port Scan", styles["body"])],
            [Paragraph("2.2 OS Detection", styles["body"])],
            [Paragraph("2.3 Insecure Protocol Detection", styles["body"])],
            [Paragraph("2.4 TLS / Weak Encryption Observations", styles["body"])],
            [Paragraph("<b>3.</b> Discovered Vulnerability Details", styles["toc_item"])],
        ]
        if include_owasp_section:
            toc_rows.extend([
                [Paragraph("<b>4.</b> OWASP Top 10:2025", styles["toc_item"])],
                [Paragraph("4.1 A01:2025 - Broken Access Control", styles["body"])],
                [Paragraph("4.2 A02:2025 - Security Misconfiguration", styles["body"])],
                [Paragraph("4.3 A03:2025 - Software Supply Chain Failures", styles["body"])],
                [Paragraph("4.4 A04:2025 - Cryptographic Failures", styles["body"])],
                [Paragraph("4.5 A05:2025 - Injection", styles["body"])],
                [Paragraph("4.6 A06:2025 - Insecure Design", styles["body"])],
                [Paragraph("4.7 A07:2025 - Authentication Failures", styles["body"])],
                [Paragraph("4.8 A08:2025 - Software or Data Integrity Failures", styles["body"])],
                [Paragraph("4.9 A09:2025 - Security Logging and Alerting Failures", styles["body"])],
                [Paragraph("4.10 A10:2025 - Mishandling of Exceptional Conditions", styles["body"])],
                [Paragraph("<b>5.</b> Conclusion", styles["toc_item"])],
            ])
        else:
            toc_rows.append([Paragraph("<b>4.</b> Conclusion", styles["toc_item"])])
        toc_content = [
            _section_banner("Table of Contents", styles),
            Spacer(1, 24),
            Table(toc_rows, colWidths=[doc.width], style=TableStyle([
                ("LEFTPADDING", (0, 0), (-1, -1), 16),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ])),
        ]
        toc_page = Table([[toc_content]], colWidths=[doc.width], rowHeights=[doc.height - 44])
        toc_page.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        story.append(toc_page)

        story.extend([
            _section_banner("1. Executive Summary", styles),
            Spacer(1, 18),
            Paragraph("Overview", styles["subsection"]),
            _safe_paragraph(
                f"This assessment summarizes the VAPT scan results for {_safe_text(payload.get('input'))}. "
                f"The report content is generated dynamically from the latest backend scan payload so values, hosts, ports, "
                f"services, and vulnerabilities change with each target.",
                styles["body"],
            ),
            Spacer(1, 8),
            _safe_paragraph(
                "The objective of this exercise is to provide an actionable security baseline: identify exposed assets, "
                "highlight weak configurations, and prioritize vulnerabilities based on operational impact and risk severity.",
                styles["body"],
            ),
            Spacer(1, 8),
            _safe_paragraph(
                product_summary.get("narrative") or "The report includes prioritized risk context, proof artifacts, and remediation guidance for engineering follow-up.",
                styles["body"],
            ),
            Spacer(1, 10),
            Paragraph("1.1 Scope of Testing", styles["section_title"]),
            _safe_paragraph(
                "The scope covered live host discovery, port identification, service detection, operating-system inference, "
                "insecure protocol review, TLS observations, service-side vulnerability mapping, and OWASP Top 10 web application checks "
                "for applicable domain-based targets.",
                styles["body"],
            ),
            Spacer(1, 8),
            _safe_paragraph(
                "The assessment was conducted as a non-disruptive security validation. It focused on externally observable services "
                "and response behavior, without attempting destructive exploitation or changes to target availability.",
                styles["body"],
            ),
            Spacer(1, 8),
            _safe_paragraph(
                "Results should be interpreted as a point-in-time view. Environmental changes such as patching, firewall updates, "
                "new services, or infrastructure drift can modify exposure and should be revalidated through periodic rescans.",
                styles["body"],
            ),
            Spacer(1, 12),
        ])
        metric_values = [
            html.escape(_safe_text(payload.get("input"))),
            str(payload.get("active_hosts", 0)),
            str(summary.get("total_vulnerabilities", 0)),
            str(owasp_summary.get("total_findings", 0)),
        ]
        metrics = Table([
            [Paragraph("Target", styles["metric_label"]), Paragraph("Active Hosts", styles["metric_label"]), Paragraph("Total Vulns", styles["metric_label"]), Paragraph("OWASP Findings", styles["metric_label"])],
            [Paragraph(metric_values[0], styles["metric_value"]), Paragraph(metric_values[1], styles["metric_value"]), Paragraph(metric_values[2], styles["metric_value"]), Paragraph(metric_values[3], styles["metric_value"])],
        ], colWidths=[doc.width / 4.0] * 4)
        metrics.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 9),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
        ]))
        story.extend([
            metrics,
            Spacer(1, 18),
            Paragraph("1.2 Graphical Summary", styles["section_title"]),
            _safe_paragraph(
                "The severity distribution below is derived from the vulnerabilities currently present in the backend response. "
                "This gives a quick operational snapshot of remediation priority.",
                styles["body"],
            ),
            Spacer(1, 10),
        ])
        graphical_summary_chart = _severity_graph_summary(severity_totals, styles, doc.width)
        if include_owasp_section:
            half_width = (doc.width - 18) / 2
            graphical_summary_chart = Table([[
                _severity_graph_summary(severity_totals, styles, half_width),
                Spacer(18, 1),
                _owasp_graph_summary(owasp_results, styles, half_width),
            ]], colWidths=[half_width, 18, half_width])
            graphical_summary_chart.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]))
        story.extend([
            graphical_summary_chart,
            Spacer(1, 16),
            Paragraph("Top 5 Risks To Fix Now", styles["section_title"]),
        ])
        top_risks = product_summary.get("top_risks") or []
        top_risk_rows = [["Risk", "Severity", "Confidence", "Why It Matters"]]
        for item in top_risks:
            top_risk_rows.append([
                _safe_paragraph(item.get("title"), styles["body_small"]),
                _safe_paragraph(item.get("severity"), styles["body_small"]),
                _safe_paragraph(f"{item.get('confidence')} ({item.get('confidence_score', 'N/A')})", styles["body_small"]),
                _safe_paragraph(item.get("business_impact"), styles["body_small"]),
            ])
        if top_risks:
            story.append(_styled_table(top_risk_rows, [170, 65, 90, 185]))
        else:
            story.append(_safe_paragraph("No prioritized risks were generated for this scan payload.", styles["body"]))
        story.extend([
            PageBreak(),
            Paragraph("1.3 List of Vulnerabilities", styles["section_title"]),
            Spacer(1, 20),
        ])
        vuln_overview_rows = [["#", "Vulnerability", "Severity", "CVSS Score", "Status"]]
        if vulnerabilities:
            for index, vulnerability in enumerate(vulnerabilities, start=1):
                vuln_overview_rows.append([
                    str(index),
                    _safe_paragraph(vulnerability["title"], styles["body_small"]),
                    _safe_paragraph(vulnerability["severity"], styles["body_small"]),
                    _safe_paragraph(str(vulnerability["cvss_score"] if vulnerability["cvss_score"] is not None else "N/A"), styles["body_small"]),
                    _safe_paragraph(vulnerability["status"], styles["body_small"]),
                ])
        else:
            vuln_overview_rows.append(["1", "No vulnerabilities discovered", "-", "-", "Closed"])
        severity_breakdown_rows = [
            ["Vulnerability Severity", "No. of Vulnerability found"],
            ["Critical", str(severity_totals.get("CRITICAL", 0))],
            ["High", str(severity_totals.get("HIGH", 0))],
            ["Medium", str(severity_totals.get("MEDIUM", 0))],
            ["Low", str(severity_totals.get("LOW", 0))],
        ]
        severity_breakdown = Table(severity_breakdown_rows, colWidths=[doc.width * 0.34, doc.width * 0.34], hAlign="CENTER")
        severity_breakdown.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 10),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("BACKGROUND", (0, 1), (0, 1), colors.HexColor("#e86334")),
            ("BACKGROUND", (0, 2), (0, 2), colors.HexColor("#f28b45")),
            ("BACKGROUND", (0, 3), (0, 3), colors.HexColor("#e3b827")),
            ("BACKGROUND", (0, 4), (0, 4), colors.HexColor("#79cc4f")),
            ("TEXTCOLOR", (0, 1), (0, 4), colors.white),
            ("FONTNAME", (0, 1), (0, 4), "Helvetica-Bold"),
            ("ALIGN", (0, 1), (0, 4), "CENTER"),
            ("BACKGROUND", (1, 1), (1, 4), colors.white),
            ("TEXTCOLOR", (1, 1), (1, 4), colors.HexColor("#111827")),
            ("ALIGN", (1, 1), (1, 4), "CENTER"),
            ("FONTNAME", (1, 1), (1, 4), "Helvetica"),
            ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#cbd5e1")),
            ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#cbd5e1")),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ]))
        vuln_overview_table = _styled_table(vuln_overview_rows, [28, 250, 70, 70, 70])
        if vulnerabilities:
            for row_index, vulnerability in enumerate(vulnerabilities, start=1):
                sev = str(vulnerability.get("severity", "")).upper()
                sev_color = {
                    "CRITICAL": colors.HexColor("#c62828"),
                    "HIGH": colors.HexColor("#ef6c00"),
                    "MEDIUM": colors.HexColor("#f59e0b"),
                    "LOW": colors.HexColor("#43a047"),
                }.get(sev, colors.HexColor("#111827"))
                vuln_overview_table.setStyle(TableStyle([
                    ("TEXTCOLOR", (2, row_index), (2, row_index), sev_color),
                    ("FONTNAME", (2, row_index), (2, row_index), "Helvetica-Bold"),
                ]))

        story.extend([
            vuln_overview_table,
            Spacer(1, 72),
            severity_breakdown,
            Spacer(1, 16),
            PageBreak(),
        ])

        story.extend([
            _section_banner("2. Technical Findings", styles),
            Spacer(1, 16),
            Paragraph("2.1 Port Scan", styles["section_title"]),
        ])
        port_rows = [["Host", "Port", "Protocol", "Service", "Product", "Version"]]
        for row in _collect_port_rows(payload):
            port_rows.append([_safe_paragraph(cell, styles["body_small"]) for cell in row])
        story.extend([
            _styled_table(port_rows, [110, 38, 56, 74, 128, 104]),
            Spacer(1, 16),
            Paragraph("2.2 OS Detection", styles["section_title"]),
        ])
        os_rows = [["Host", "OS Name"]]
        for row in _collect_os_rows(payload):
            os_rows.append([_safe_paragraph(cell, styles["body_small"]) for cell in row[:2]])
        story.extend([
            _styled_table(os_rows, [180, 330]),
            Spacer(1, 16),
            Paragraph("2.3 Insecure Protocol Detection", styles["section_title"]),
        ])
        insecure_rows = [["Host", "Port", "Protocol", "Observation"]]
        for row in _collect_protocol_rows(payload, "insecure_protocols"):
            insecure_rows.append([_safe_paragraph(cell, styles["body_small"]) for cell in row])
        story.extend([
            _styled_table(insecure_rows, [110, 42, 90, 268]),
            Spacer(1, 16),
            Paragraph("2.4 TLS / Weak Encryption Observations", styles["section_title"]),
        ])
        tls_rows = [["Host", "Port", "Version / Type", "Observation"]]
        for row in _collect_protocol_rows(payload, "tls_issues"):
            tls_rows.append([_safe_paragraph(cell, styles["body_small"]) for cell in row])
        story.extend([
            _styled_table(tls_rows, [110, 42, 100, 258]),
            PageBreak(),
        ])

        if vulnerabilities:
            story.extend([
                _section_banner("3. Discovered Vulnerability Details", styles),
                Spacer(1, 16),
            ])
            for index, vulnerability in enumerate(vulnerabilities, start=1):
                story.extend([
                    Paragraph(f"Vulnerability #{index}", styles["section_title"]),
                    Paragraph(html.escape(vulnerability["title"]), styles["subsection"]),
                    Spacer(1, 8),
                ])
                detail_meta = [[
                    Paragraph(f"<b>Severity</b><br/>{html.escape(vulnerability['severity'])}", styles["body"]),
                    Paragraph(f"<b>Status</b><br/>{html.escape(vulnerability['status'])}", styles["body"]),
                    Paragraph(f"<b>CVSS Score</b><br/>{html.escape(str(vulnerability['cvss_score'] if vulnerability['cvss_score'] is not None else 'N/A'))}", styles["body"]),
                    Paragraph(f"<b>Confidence</b><br/>{html.escape(str(vulnerability.get('confidence') or 'unknown'))} ({html.escape(str(vulnerability.get('confidence_score') if vulnerability.get('confidence_score') is not None else 'N/A'))})", styles["body"]),
                ]]
                meta_table = Table(detail_meta, colWidths=[120, 120, 120, 120])
                meta_table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
                    ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]))
                story.extend([
                    meta_table,
                    Spacer(1, 12),
                    _safe_paragraph(f"Affected URL / Service: {vulnerability['host']} ({vulnerability['ip']}) - {vulnerability['service']}:{vulnerability['port']}", styles["body"]),
                    Spacer(1, 8),
                    Paragraph("Details of Vulnerability", styles["section_title"]),
                    _safe_paragraph(vulnerability["description"], styles["body"]),
                    Spacer(1, 8),
                    Paragraph("Impact", styles["section_title"]),
                    _safe_paragraph(_impact_text(vulnerability), styles["body"]),
                    Spacer(1, 8),
                    Paragraph("Validation & Proof", styles["section_title"]),
                    _safe_paragraph(
                        f"Validation state: {vulnerability.get('validation_state', 'unknown')}. {_proof_summary(vulnerability)}",
                        styles["body"],
                    ),
                    Spacer(1, 8),
                    Paragraph("Remediation", styles["section_title"]),
                    _safe_paragraph(vulnerability["remediation"], styles["body"]),
                    Spacer(1, 8),
                    Paragraph("Suggested Fixes", styles["section_title"]),
                    _safe_paragraph(_suggested_fixes_text(vulnerability), styles["body"]),
                    Spacer(1, 8),
                    Paragraph("Compliance Mapping", styles["section_title"]),
                    _safe_paragraph(", ".join(vulnerability.get("compliance_mapping") or ["No compliance mapping available."]), styles["body"]),
                    Spacer(1, 8),
                    Paragraph("Additional References", styles["section_title"]),
                ])
                ref_rows = [["Reference"]]
                for reference in _additional_references(vulnerability):
                    ref_rows.append([_safe_paragraph(reference, styles["body_small"])])
                story.append(_styled_table(ref_rows, [doc.width]))
                if index != len(vulnerabilities):
                    story.append(PageBreak())
        else:
            story.extend([
                _section_banner("3. Discovered Vulnerability Details", styles),
                Spacer(1, 16),
                _safe_paragraph("No vulnerabilities were detected in the selected scan payload.", styles["body"]),
            ])

        if include_owasp_section:
            story.extend([
                PageBreak(),
                _section_banner("4. OWASP Top 10:2025", styles),
                Spacer(1, 16),
            ])
            owasp_summary_rows = [
                ["Total Categories", "Categories with Findings", "Total Findings", "Target URL"],
                [
                    str(owasp_summary.get("total_categories", 0)),
                    str(owasp_summary.get("categories_with_findings", 0)),
                    str(owasp_summary.get("total_findings", 0)),
                    _safe_paragraph(payload.get("owasp_top_10", {}).get("normalized_url") or "N/A", styles["body_small"]),
                ],
            ]
            story.extend([
                Paragraph("OWASP Summary", styles["section_title"]),
                _styled_table(owasp_summary_rows, [90, 120, 90, 230]),
                Spacer(1, 16),
            ])
            if owasp_results:
                category_rows = [["ID", "Category", "Status", "Severity", "Findings"]]
                for item in owasp_results:
                    category_rows.append([
                        item.get("short") or item.get("id") or "-",
                        _safe_paragraph(item.get("title") or "-", styles["body_small"]),
                        _safe_paragraph((item.get("status") or "completed").title(), styles["body_small"]),
                        _safe_paragraph(item.get("severity") or "-", styles["body_small"]),
                        _safe_paragraph(str(item.get("findings_count", 0)), styles["body_small"]),
                    ])
                story.extend([
                    Paragraph("OWASP Category Overview", styles["section_title"]),
                    _styled_table(category_rows, [50, 250, 80, 80, 80]),
                    Spacer(1, 18),
                ])
                for item in owasp_results:
                    story.extend([
                        Paragraph(f"{item.get('id', 'OWASP')} - {item.get('title', 'Category')}", styles["subsection"]),
                    ])
                    findings = item.get("findings") or []
                    if findings:
                        finding_rows = [["#", "Finding", "Details", "URL"]]
                        for index, finding in enumerate(findings, start=1):
                            finding_rows.append([
                                str(index),
                                _safe_paragraph(finding.get("title") or "Finding", styles["body_small"]),
                                _safe_paragraph(finding.get("description") or finding.get("evidence") or "No details provided.", styles["body_small"]),
                                _safe_paragraph(finding.get("url") or "N/A", styles["body_small"]),
                            ])
                        story.extend([
                            _styled_table(finding_rows, [24, 156, 210, 120]),
                            Spacer(1, 12),
                        ])
                    else:
                        story.extend([
                            _safe_paragraph("No obvious findings were detected for this OWASP category during the current scan window.", styles["body"]),
                            Spacer(1, 10),
                        ])
            else:
                story.append(_safe_paragraph("The OWASP checks completed without recording category-level findings.", styles["body"]))

        total_vulns = int(summary.get("total_vulnerabilities", 0) or 0)
        critical_count = int(summary.get("critical_risk", 0) or 0)
        high_count = int(summary.get("high_risk", 0) or 0)
        active_hosts = int(payload.get("active_hosts", 0) or 0)
        target_label = _safe_text(payload.get("input"))
        owasp_findings = int(owasp_summary.get("total_findings", 0) or 0)
        if total_vulns == 0:
            conclusion_text = (
                f"The assessment for {target_label} identified {active_hosts} active host(s), no confirmed network-service vulnerabilities, "
                f"and {owasp_findings} OWASP web finding(s) in the current scan window. This indicates a comparatively stronger baseline posture "
                "for exposed services, but continuous monitoring, application hardening, and periodic re-validation remain essential."
            )
        else:
            conclusion_text = (
                f"The assessment for {target_label} identified {active_hosts} active host(s) with {total_vulns} total vulnerability finding(s), "
                f"including {critical_count} critical and {high_count} high severity issue(s), alongside {owasp_findings} OWASP web finding(s). "
                "Immediate remediation should prioritize critical and high-risk exposures, followed by application-layer weaknesses, through a tracked closure plan and verification rescans."
            )
        story.extend([
            PageBreak(),
            _section_banner("5. Conclusion" if include_owasp_section else "4. Conclusion", styles),
            Spacer(1, 18),
            _safe_paragraph(conclusion_text, styles["body"]),
        ])

        doc.build(story, onFirstPage=_draw_page_footer, onLaterPages=_draw_page_footer)
        return

    # fallback minimal valid PDF writer (no reportlab)
    lines = ["VAPT Scan Report", f"Generated: {datetime.utcnow().isoformat()} UTC", "", f"Target: {payload.get('input','N/A')}"]
    lines += [f"Host: {asset.get('hostname',asset.get('domain','Unknown'))}, IP: {asset.get('ip','Unknown')}, OS: {asset.get('os',asset.get('os_name','Unknown'))}, Vendor: {asset.get('vendor','Unknown')}" for asset in payload.get('assets', [])]
    text = "\n".join(lines)
    text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    content_stream = ["BT", "/F1 12 Tf", "50 760 Td"]
    y = 760
    for line in text.split('\n'):
        escaped = line
        content_stream.append(f"({escaped}) Tj")
        y -= 14
        content_stream.append("0 -14 Td")
    content_stream.append("ET")
    stream_data = "\n".join(content_stream).encode('latin-1', 'replace')

    # Minimal PDF object construction
    objs = []
    objs.append((1, b"<< /Type /Catalog /Pages 2 0 R >>\n"))
    objs.append((2, b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>\n"))
    page_obj = (b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\n")
    objs.append((3, page_obj))
    objs.append((4, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\n"))
    contents = b"".join([b"<< /Length %d >>\nstream\n" % len(stream_data), stream_data, b"\nendstream\n"])
    objs.append((5, contents))

    with open(output_path, 'wb') as f:
        f.write(b"%PDF-1.4\n")
        xref = []
        for num, body in objs:
            xref.append(f.tell())
            f.write(f"{num} 0 obj\n".encode())
            f.write(body)
            f.write(b"endobj\n")

        start_xref = f.tell()
        f.write(b"xref\n")
        f.write(f"0 {len(objs)+1}\n".encode())
        f.write(b"0000000000 65535 f \n")
        for pos in xref:
            f.write(f"{pos:010d} 00000 n \n".encode())
        f.write(b"trailer\n")
        f.write(b"<< /Size ")
        f.write(f"{len(objs)+1}".encode())
        f.write(b" /Root 1 0 R >>\nstartxref\n")
        f.write(f"{start_xref}\n".encode())
        f.write(b"%%EOF\n")


def _generate_docx_report(payload: dict, output_path: str):
    if Document is not None:
        doc = Document()
        doc.add_heading("VAPT Scan Report", level=1)
        doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()} UTC")
        doc.add_paragraph(f"Target: {payload.get('input', 'N/A')}")

        summary = payload.get('vulnerability_summary', {})
        product_summary = payload.get("product_summary", {})
        owasp_summary = _collect_owasp_summary(payload)
        doc.add_paragraph(f"Active hosts: {payload.get('active_hosts', 0)}")
        doc.add_paragraph(f"Total vulnerabilities: {summary.get('total_vulnerabilities', 0)}")
        doc.add_paragraph(f"Critical risk: {summary.get('critical_risk', 0)}")
        doc.add_paragraph(f"OWASP findings: {owasp_summary.get('total_findings', 0)}")
        if product_summary.get("narrative"):
            doc.add_paragraph(product_summary["narrative"])
        if product_summary.get("top_risks"):
            doc.add_paragraph("Top risks to fix now:")
            for item in product_summary["top_risks"]:
                doc.add_paragraph(
                    f"{item.get('severity')} - {item.get('title')} [{item.get('confidence')} {item.get('confidence_score')}]",
                    style='List Bullet',
                )

        assets = payload.get('assets', [])
        if assets:
            table = doc.add_table(rows=1, cols=6)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            headings = ['Host', 'IP', 'OS', 'Vendor', 'Open ports', 'Vulns']
            for i, heading in enumerate(headings):
                hdr_cells[i].text = heading

            for asset in assets:
                row_cells = table.add_row().cells
                row_cells[0].text = str(asset.get('hostname', asset.get('domain', 'Unknown')))
                row_cells[1].text = str(asset.get('ip', 'Unknown'))
                row_cells[2].text = str(asset.get('os', asset.get('os_name', 'Unknown')))
                row_cells[3].text = str(asset.get('vendor', 'Unknown'))
                open_ports = ', '.join(str(p.get('port')) for p in asset.get('open_ports', []) if isinstance(p, dict)) or '-'
                row_cells[4].text = open_ports
                row_cells[5].text = str(len(asset.get('vulnerabilities', [])))

        doc.add_paragraph('\nVulnerability detail section')
        for asset in assets:
            vulns = asset.get('vulnerabilities', [])
            if not vulns:
                continue
            doc.add_paragraph(f"{asset.get('hostname', asset.get('domain', 'Unknown'))}: {len(vulns)} vuln(s)", style='List Bullet')

        owasp_results = _collect_owasp_results(payload)
        if owasp_results:
            doc.add_paragraph("\nOWASP Top 10")
            for item in owasp_results:
                doc.add_paragraph(f"{item.get('id', 'OWASP')} {item.get('title', 'Category')}: {item.get('findings_count', 0)} finding(s)", style='List Bullet')

        doc.save(output_path)
        return

    # fallback minimal docx file (raw package)
    import zipfile

    def xml_escape(text):
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&apos;')

    body = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">',
        '<w:body>',
        '<w:p><w:r><w:t>VAPT Scan Report</w:t></w:r></w:p>',
        f'<w:p><w:r><w:t>Generated: {xml_escape(datetime.utcnow().isoformat())} UTC</w:t></w:r></w:p>',
        f'<w:p><w:r><w:t>Target: {xml_escape(str(payload.get("input", "N/A")))}</w:t></w:r></w:p>',
    ]

    for asset in payload.get('assets', []):
        host = xml_escape(str(asset.get('hostname', asset.get('domain', 'Unknown'))))
        line = f"{host} - IP: {xml_escape(str(asset.get('ip', 'Unknown')))} - OS: {xml_escape(str(asset.get('os', asset.get('os_name', 'Unknown'))))}"
        body.append(f'<w:p><w:r><w:t>{line}</w:t></w:r></w:p>')

    body.append('</w:body></w:document>')
    document_xml = '\n'.join(body).encode('utf-8')

    package_parts = {
        '[Content_Types].xml': b'<?xml version="1.0" encoding="UTF-8"?>' +
                             b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">' +
                             b'<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>' +
                             b'<Default Extension="xml" ContentType="application/xml"/>' +
                             b'<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>' +
                             b'</Types>',
        '_rels/.rels': b'<?xml version="1.0" encoding="UTF-8"?>' +
                       b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">' +
                       b'<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>' +
                       b'</Relationships>',
        'word/_rels/document.xml.rels': b'<?xml version="1.0" encoding="UTF-8"?>' +
                                       b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>',
        'word/document.xml': document_xml,
    }

    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for name, content in package_parts.items():
            zf.writestr(name, content)

        doc.add_paragraph("")

        for line in json.dumps(payload, indent=2).splitlines():
            doc.add_paragraph(line)

        doc.save(output_path)
        return

    # Fallback minimal docx writer
    import zipfile

    def xml_escape(s):
        return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                .replace("\"", "&quot;").replace("'", "&apos;"))

    body_lines = ["<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>",
                  "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">",
                  "<w:body>"]
    body_lines.append("<w:p><w:r><w:t>VAPT Scan Report</w:t></w:r></w:p>")
    body_lines.append(f"<w:p><w:r><w:t>Generated: {datetime.utcnow().isoformat()} UTC</w:t></w:r></w:p>")
    body_lines.append("<w:p><w:r><w:t/></w:r></w:p>")

    for line in json.dumps(payload, indent=2).splitlines():
        body_lines.append(f"<w:p><w:r><w:t>{xml_escape(line)}</w:t></w:r></w:p>")

    body_lines.append("</w:body></w:document>")
    document_xml = "".join(body_lines).encode("utf-8")

    package_parts = {
        "[Content_Types].xml": b"<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
                            b"<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
                            b"<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
                            b"<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
                            b"<Override PartName=\"/word/document.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
                            b"</Types>",
        "_rels/.rels": b"<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
                        b"<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
                        b"<Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"word/document.xml\"/>"
                        b"</Relationships>",
        "word/_rels/document.xml.rels": b"<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
                                       b"<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
                                       b"</Relationships>",
        "word/document.xml": document_xml,
    }

    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, content in package_parts.items():
            zf.writestr(name, content)


def save_report(target: str, payload: dict):
    """Save scan data in JSON and TXT report files."""

    os.makedirs("reports", exist_ok=True)
    safe_target = target.replace("/", "-").replace("\\", "-").replace(":", "-").replace(" ", "_")
    tstamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"reports/scan_{safe_target}_{tstamp}"

    json_path = f"{base_name}.json"
    txt_path = f"{base_name}.txt"

    with open(json_path, "w", encoding="utf-8") as jf:
        json.dump(payload, jf, indent=2)

    with open(txt_path, "w", encoding="utf-8") as tf:
        tf.write("Scan Report\n")
        tf.write("============\n\n")
        tf.write(json.dumps(payload, indent=2))

    return {"json": json_path, "txt": txt_path}


def _scan_target_type(target: str) -> str:
    value = str(target or "").strip()
    if "," in value or "/" in value or "-" in value:
        return "Range"
    return "Domain" if is_domain_input(value) else "IP"


def _persist_scan_to_db(payload: dict):
    db = SessionLocal()
    try:
        summary = payload.get("vulnerability_summary", {}) or {}
        report_files = payload.get("report_files", {}) or {}
        assets = payload.get("assets", []) or []
        scan = ScanModel(
            target=payload.get("input"),
            target_type=_scan_target_type(payload.get("input")),
            status="completed",
            active_hosts=int(payload.get("active_hosts", 0) or 0),
            total_hosts=int(payload.get("total_targets", 0) or len(assets)),
            total_vulnerabilities=int(summary.get("total_vulnerabilities", 0) or 0),
            critical_risk=int(summary.get("critical_risk", 0) or 0),
            report_json_path=report_files.get("json"),
            report_txt_path=report_files.get("txt"),
            report_pdf_path=report_files.get("pdf"),
        )
        db.add(scan)
        db.flush()

        for asset in assets:
            asset_row = AssetModel(
                scan_id=scan.id,
                target_input=payload.get("input"),
                ip=asset.get("ip"),
                hostname=asset.get("hostname"),
                domain=asset.get("domain"),
                mac=asset.get("mac"),
                vendor=asset.get("vendor"),
                device_type=asset.get("device_type"),
                os_name=asset.get("os") or asset.get("os_name"),
                status=asset.get("status") or "UP",
                country=asset.get("country"),
            )
            db.add(asset_row)
            db.flush()

            for port_info in asset.get("open_ports", []) or []:
                db.add(PortModel(
                    asset_id=asset_row.id,
                    port=port_info.get("port"),
                    protocol=port_info.get("protocol"),
                    service=port_info.get("service"),
                    product=port_info.get("product"),
                    version=port_info.get("version"),
                    state=port_info.get("state") or "open",
                ))

            for vuln in asset.get("vulnerabilities", []) or []:
                db.add(VulnerabilityModel(
                    asset_id=asset_row.id,
                    cve_id=vuln.get("cve") or vuln.get("cve_id"),
                    title=vuln.get("title"),
                    description=vuln.get("description"),
                    severity=vuln.get("severity"),
                    cvss_score=vuln.get("cvss_score"),
                    product=vuln.get("product"),
                    version=vuln.get("version"),
                    remediation=vuln.get("remediation"),
                    status=vuln.get("status") or "open",
                ))

            for item in asset.get("insecure_protocols", []) or []:
                db.add(InsecureProtocol(
                    asset_id=asset_row.id,
                    port=item.get("port"),
                    protocol=item.get("protocol"),
                    message=item.get("message") or item.get("msg"),
                ))

            for item in asset.get("tls_issues", []) or []:
                db.add(TLSIssue(
                    asset_id=asset_row.id,
                    port=item.get("port"),
                    tls_version=item.get("tls_version") or item.get("protocol"),
                    message=item.get("message") or item.get("msg"),
                ))

        owasp = payload.get("owasp_top_10") or {}
        for category in owasp.get("results", []) or []:
            category_row = OWASPResult(
                scan_id=scan.id,
                category_id=category.get("id") or category.get("short"),
                category_name=category.get("title"),
                severity=category.get("severity"),
                findings_count=int(category.get("findings_count", 0) or 0),
                status=category.get("status") or "completed",
            )
            db.add(category_row)
            db.flush()

            for finding in category.get("findings", []) or []:
                db.add(OWASPFinding(
                    owasp_result_id=category_row.id,
                    title=finding.get("title"),
                    description=finding.get("description"),
                    severity=finding.get("severity"),
                    url=finding.get("url"),
                    evidence=finding.get("evidence"),
                ))

        db.commit()
        return scan.id
    except Exception:
        db.rollback()
        raise
    finally:
        db.close()


def parse_targets(target: str):
    targets = []

    # Multiple inputs (comma-separated)
    if "," in target:
        parts = [t.strip() for t in target.split(",")]
        for part in parts:
            targets.extend(parse_targets(part))
        return targets

    # IP range (start-end)
    if "-" in target and target.count("-") == 1:
        start_ip, end_ip = target.split("-")
        try:
            start = ipaddress.IPv4Address(start_ip.strip())
            end = ipaddress.IPv4Address(end_ip.strip())
        except ValueError:
            # Not an IP range (e.g., domain names may contain a dash)
            pass
        else:
            for ip_int in range(int(start), int(end) + 1):
                targets.append(str(ipaddress.IPv4Address(ip_int)))
            return targets

    # CIDR subnet
    try:
        network = ipaddress.ip_network(target, strict=False)
        return [str(ip) for ip in network.hosts()]
    except ValueError:
        pass

    # Single IP
    try:
        ipaddress.ip_address(target)
        return [target]
    except ValueError:
        # Try resolving as a domain name
        resolved = resolve_domain(target)
        if resolved:
            return resolved

        raise ValueError(f"Invalid target format: {target}")


def validate_target(target: str) -> tuple[bool, str]:
    """
    Validate target format before scanning.
    Returns (is_valid, error_message)
    """
    if not target or len(target.strip()) == 0:
        return False, "Target cannot be empty"

    # Split by comma for multiple targets
    parts = [p.strip() for p in target.split(",")]

    for part in parts:
        # Try to parse as CIDR
        try:
            ipaddress.ip_network(part, strict=False)
            continue
        except ValueError:
            pass

        # Try to parse as IP address
        try:
            ipaddress.ip_address(part)
            continue
        except ValueError:
            pass

        # Try to parse as IP range (192.168.1.1-192.168.1.10)
        if "-" in part and part.count("-") == 1:
            try:
                start_ip, end_ip = part.split("-")
                ipaddress.IPv4Address(start_ip.strip())
                ipaddress.IPv4Address(end_ip.strip())
                continue
            except ValueError:
                pass

        # Try to parse as domain
        domain_pattern = r"^([a-zA-Z0-9]([a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}$"
        if re.match(domain_pattern, part):
            resolved = resolve_domain(part)
            if not resolved:
                return False, f"Domain '{part}' could not be resolved. Please check spelling or DNS settings."
            continue

        # Accept single-label hostnames for local network names like 'm' or 'printer1'
        local_host_pattern = r"^[a-zA-Z0-9]([a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?$"
        if re.match(local_host_pattern, part):
            continue

        return False, f"Invalid target format: '{part}'"

    return True, ""


@router.get("/validate_target")
def validate_target_endpoint(target: str = Query(...)):
    is_valid, error_msg = validate_target(target)
    return {
        "valid": is_valid,
        "error": error_msg
    }


def _execute_scan(target: str, job_id: str | None = None, options: dict | None = None):
    is_valid, error_msg = validate_target(target)
    if not is_valid:
        return {"error": error_msg, "success": False}

    options = options or {}
    auth_context = options.get("auth") or {}
    include_surface_discovery = bool(options.get("include_surface_discovery", True))
    include_api_security = bool(options.get("include_api_security", True))
    api_document = options.get("api_document")
    api_document_format = options.get("api_document_format")
    api_base_url = options.get("api_base_url")

    results = []
    include_domain = is_domain_input(target)

    if job_id:
        _job_update(job_id, status="running", progress=3, stage_index=0, stage_label="Host Discovery")
        _job_log(job_id, f"Starting VAPT scan for target: {target}")

    try:
        ip_list = list(set(parse_targets(target)))
    except ValueError as e:
        return {"error": str(e)}

    if job_id and not _job_wait_if_paused(job_id):
        return {"error": "Scan cancelled"}

    if job_id:
        _job_log(job_id, f"Parsed {len(ip_list)} target(s). Running host discovery.")
    active_hosts = threaded_ping_sweep(ip_list)

    if not active_hosts:
        if job_id:
            _job_log(job_id, "No ICMP replies; falling back to direct target scanning.")
        active_hosts = ip_list

    if job_id:
        _job_update(job_id, progress=22, stage_index=1, stage_label="Port Scanning (1-65535)")
        _job_log(job_id, f"Host discovery complete. {len(active_hosts)} host(s) queued for scanning.")

    total_hosts = max(len(active_hosts), 1)

    for idx, ip in enumerate(active_hosts):
        if job_id and not _job_wait_if_paused(job_id):
            return {"error": "Scan cancelled"}

        ports = threaded_port_scan(ip)
        if job_id:
            _job_update(job_id, progress=min(42, 22 + int(((idx + 1) / total_hosts) * 20)), stage_index=1, stage_label="Port Scanning (1-65535)")
            _job_log(job_id, f"Port scan finished for {ip}: {len(ports)} open port(s) detected.")

        if job_id and not _job_wait_if_paused(job_id):
            return {"error": "Scan cancelled"}

        services = detect_services(ip, ports)
        if not services and ports:
            services = [
                {"port": p, "protocol": "tcp", "service": "unknown", "product": "", "version": ""}
                for p in sorted(set(ports))
            ]

        hostname = get_hostname(ip) or ip
        asset_domain = None
        if include_domain:
            asset_domain = target
            hostname = target
        elif hostname and hostname != ip:
            asset_domain = hostname

        vendor = "Unknown"
        if not settings.FAST_SCAN_MODE:
            mac = get_mac(ip)
            vendor = get_vendor(mac)
        if include_domain and (vendor == "Unknown" or not vendor):
            vendor = target

        if job_id:
            _job_update(job_id, progress=min(62, 45 + int(((idx + 1) / total_hosts) * 17)), stage_index=2, stage_label="Service Detection")
            _job_log(job_id, f"Service detection finished for {ip}.")

        if job_id and not _job_wait_if_paused(job_id):
            return {"error": "Scan cancelled"}

        device = classify_device(ip, ports, vendor)
        os_details = detect_os_details(ip, services=services, vendor=vendor, device_type=device, hostname=hostname)
        insecure_protocols = detect_insecure_protocols(services)

        tls_findings = []
        tls_ports = set([s.get("port") for s in services if s.get("port") in (443, 8443)])
        tls_ports.update([p for p in ports if p in (443, 8443)])
        for tport in sorted(tls_ports):
            tls_findings.extend(analyze_tls(ip, tport))

        vuls = assess_vulnerabilities(services)

        asset = {
            "ip": ip,
            "hostname": hostname,
            "vendor": vendor,
            "os": os_details["name"],
            "open_ports": services,
            "device_type": device,
            "insecure_protocols": insecure_protocols,
            "tls_issues": tls_findings,
            "vulnerabilities": vuls
        }

        if include_domain or asset_domain:
            asset["domain"] = asset_domain
        if include_domain:
            asset["resolved_ip"] = ip

        results.append(asset)

        if job_id:
            _job_update(job_id, progress=min(85, 65 + int(((idx + 1) / total_hosts) * 20)), stage_index=3, stage_label="Vulnerability Assessment")
            _job_log(job_id, f"Vulnerability assessment finished for {ip}: {len(vuls)} finding(s).")

    all_vulnerabilities = [v for asset in results for v in asset.get("vulnerabilities", [])]
    critical_risks = [v for v in all_vulnerabilities if v.get("severity") == "CRITICAL"]
    high_risks = [v for v in all_vulnerabilities if v.get("severity") == "HIGH"]
    if job_id:
        _job_update(job_id, progress=92, stage_index=4, stage_label="Security Analysis")
        _job_log(job_id, "Running OWASP Top 10 web checks.")
    owasp_top_10 = _run_owasp_scan(target, include_domain, auth_context)
    web_surface = run_discover_web_surface(target, auth_context) if include_domain and include_surface_discovery else {
        "base_url": None,
        "pages": [],
        "scripts": [],
        "api_candidates": [],
        "hidden_routes": [],
        "auth_context_used": bool(auth_context),
    }
    api_security = {"checked_endpoints": [], "findings": [], "summary": {"checked": 0, "findings": 0}}
    if include_domain and include_api_security:
        if api_document:
            api_security = run_analyze_api_document(api_document, api_document_format, api_base_url or web_surface.get("base_url"), auth_context)
        else:
            api_security = run_assess_api_endpoints(web_surface.get("api_candidates", []), auth_context)
    product_summary = executive_summary({"assets": results})
    response = {
        "input": target,
        "total_targets": len(ip_list),
        "active_hosts": len(results),
        "assets": results,
        "owasp_top_10": owasp_top_10,
        "web_surface": web_surface,
        "api_security": api_security,
        "scan_profile": {
            "authenticated": bool(auth_context),
            "include_surface_discovery": include_surface_discovery,
            "include_api_security": include_api_security,
            "api_document_supplied": bool(api_document),
        },
        "product_summary": product_summary,
        "vulnerability_summary": {
            "total_vulnerabilities": len(all_vulnerabilities),
            "critical_risk": len(critical_risks),
            "high_risk": len(high_risks),
            "confirmed_findings": product_summary.get("confirmed_findings", 0),
            "needs_validation": product_summary.get("needs_validation", 0),
            "owasp_findings": _owasp_findings_count({"owasp_top_10": owasp_top_10}),
            "api_findings": int((api_security.get("summary") or {}).get("findings", 0) or 0),
        }
    }
    response["report_files"] = save_report(target, response)
    try:
        response["db_scan_id"] = _persist_scan_to_db(response)
    except Exception as exc:
        response["db_error"] = str(exc)

    if job_id:
        _job_update(job_id, progress=95, stage_index=4, stage_label="Security Analysis")
        _job_log(job_id, "Security analysis complete. Finalizing report data.")

    return response


@router.get("/scan")
def run_scan(target: str = Query(...)):
    return _execute_scan(target)


@router.post("/scan")
def run_scan_with_options(payload: dict = Body(...)):
    target = str(payload.get("target") or "").strip()
    if not target:
        raise HTTPException(status_code=400, detail="target is required")
    return _execute_scan(target, options=payload.get("options") or {})

    # Validate target first
    is_valid, error_msg = validate_target(target)
    if not is_valid:
        return {"error": error_msg, "success": False}

    results = []
    include_domain = is_domain_input(target)
    print(f"Received scan request for target: {target}")
    # 🔹 Step 1: Parse all input formats
    try:
        ip_list = list(set(parse_targets(target)))  # remove duplicates
    except ValueError as e:
        return {"error": str(e)}

    # 🔹 Step 2: Find active hosts (ICMP ping)
    active_hosts = threaded_ping_sweep(ip_list)

    if not active_hosts:
        # Some hosts may block ICMP, fallback to scanning the provided targets directly
        print("No ICMP replies; falling back to port scanning all parsed targets.")
        active_hosts = ip_list

    # 🔹 Step 3: Scan each active host
    for ip in active_hosts:
        ports = threaded_port_scan(ip)
        # Always run service detection; if no ports were found from the TCP/UDP probe,
        # the detector can fallback to full service scan (-p-).
        services = detect_services(ip, ports)

        # If nmap service/version scan could not detect details, preserve basic port data
        if not services and ports:
            services = [
                {
                    "port": p,
                    "protocol": "tcp",
                    "service": "unknown",
                    "product": "",
                    "version": ""
                }
                for p in sorted(set(ports))
            ]

        hostname = get_hostname(ip)

        # Ensure we always have a hostname recorded (at least IP fallback).
        if not hostname:
            hostname = ip

        # For top-level domain scan, preserve the input domain in hostname (desired UX)
        # and store the resolved IP in a dedicated field for frontend display.
        asset_domain = None
        if include_domain:
            asset_domain = target
            hostname = target
        else:
            # During IP based scan, keep reverse-resolved hostname if available
            if hostname and hostname != ip:
                asset_domain = hostname

        # Use target domain as vendor if no OUI vendor is detected for domain scanning
        mac = None
        vendor = "Unknown"
        if not settings.FAST_SCAN_MODE:
            mac = get_mac(ip)
            vendor = get_vendor(mac)

        if include_domain and (vendor == "Unknown" or not vendor):
            vendor = target

        device = classify_device(ip, ports, vendor)
        os_details = detect_os_details(
            ip,
            services=services,
            vendor=vendor,
            device_type=device,
            hostname=hostname,
        )

        print(f"[run_scan] {ip} => port probe: {ports}, services detected: {len(services)}, os: {os_details['name']}")

        insecure_protocols = detect_insecure_protocols(services)

        tls_findings = []
        tls_ports = set([s.get("port") for s in services if s.get("port") in (443, 8443)])
        tls_ports.update([p for p in ports if p in (443, 8443)])
        for tport in sorted(tls_ports):
            tls_findings.extend(analyze_tls(ip, tport))

        vuls = assess_vulnerabilities(services)

        asset = {
            "ip": ip,
            "hostname": hostname,
            "vendor": vendor,
            "os": os_details["name"],
            "open_ports": services,
            "device_type": device,
            "insecure_protocols": insecure_protocols,
            "tls_issues": tls_findings,
            "vulnerabilities": vuls
        }

        if include_domain or asset_domain:
            asset["domain"] = asset_domain
        if include_domain:
            asset["resolved_ip"] = ip

        results.append(asset)

    # Aggregate vulnerability and risk stats
    all_vulnerabilities = [v for asset in results for v in asset.get("vulnerabilities", [])]
    critical_risks = [v for v in all_vulnerabilities if v.get("severity") == "CRITICAL"]

    response = {
        "input": target,
        "total_targets": len(ip_list),
        "active_hosts": len(results),
        "assets": results,
        "vulnerability_summary": {
            "total_vulnerabilities": len(all_vulnerabilities),
            "critical_risk": len(critical_risks)
        }
    }

    report_paths = save_report(target, response)
    response["report_files"] = report_paths

    return response


def _run_scan_job(job_id: str, target: str):
    try:
        with SCAN_JOBS_LOCK:
            job = SCAN_JOBS.get(job_id) or {}
            options = job.get("options") or {}
        result = _execute_scan(target, job_id=job_id, options=options)
        with SCAN_JOBS_LOCK:
            job = SCAN_JOBS.get(job_id)
            if not job:
                return
            if job.get("cancel_requested"):
                job["status"] = "cancelled"
                job["error"] = "Scan cancelled"
            elif result.get("error"):
                job["status"] = "error"
                job["error"] = result["error"]
            else:
                job["status"] = "completed"
                job["progress"] = 100
                job["stage_index"] = 4
                job["stage_label"] = "Security Analysis"
                job["result"] = result
            job["updated_at"] = datetime.utcnow().isoformat() + "Z"
    except Exception as exc:
        with SCAN_JOBS_LOCK:
            job = SCAN_JOBS.get(job_id)
            if job:
                job["status"] = "error"
                job["error"] = str(exc)
                job["updated_at"] = datetime.utcnow().isoformat() + "Z"


@router.get("/scan/start")
def start_scan_job(target: str = Query(...)):
    is_valid, error_msg = validate_target(target)
    if not is_valid:
        return {"error": error_msg, "success": False}

    job_id = uuid.uuid4().hex
    now = datetime.utcnow().isoformat() + "Z"
    job = {
        "job_id": job_id,
        "target": target,
        "options": {},
        "status": "queued",
        "progress": 0,
        "stage_index": 0,
        "stage_label": "Host Discovery",
        "logs": [],
        "result": None,
        "error": None,
        "pause_requested": False,
        "cancel_requested": False,
        "created_at": now,
        "updated_at": now,
    }
    with SCAN_JOBS_LOCK:
        SCAN_JOBS[job_id] = job

    thread = threading.Thread(target=_run_scan_job, args=(job_id, target), daemon=True)
    thread.start()
    return _job_snapshot(job)


@router.post("/scan/start")
def start_scan_job_with_options(payload: dict = Body(...)):
    target = str(payload.get("target") or "").strip()
    if not target:
        raise HTTPException(status_code=400, detail="target is required")

    is_valid, error_msg = validate_target(target)
    if not is_valid:
        return {"error": error_msg, "success": False}

    job_id = uuid.uuid4().hex
    now = datetime.utcnow().isoformat() + "Z"
    job = {
        "job_id": job_id,
        "target": target,
        "options": payload.get("options") or {},
        "status": "queued",
        "progress": 0,
        "stage_index": 0,
        "stage_label": "Host Discovery",
        "logs": [],
        "result": None,
        "error": None,
        "pause_requested": False,
        "cancel_requested": False,
        "created_at": now,
        "updated_at": now,
    }
    with SCAN_JOBS_LOCK:
        SCAN_JOBS[job_id] = job

    thread = threading.Thread(target=_run_scan_job, args=(job_id, target), daemon=True)
    thread.start()
    return _job_snapshot(job)


@router.get("/scan/status")
def scan_job_status(job_id: str = Query(...)):
    with SCAN_JOBS_LOCK:
        job = SCAN_JOBS.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Scan job not found")
        return _job_snapshot(job)


@router.get("/scan/pause")
def pause_scan_job(job_id: str = Query(...)):
    with SCAN_JOBS_LOCK:
        job = SCAN_JOBS.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Scan job not found")
        if job["status"] in ("completed", "error", "cancelled"):
            return _job_snapshot(job)
        job["pause_requested"] = True
        job["status"] = "paused"
        job["updated_at"] = datetime.utcnow().isoformat() + "Z"
    _job_log(job_id, "Scan paused by user.")
    with SCAN_JOBS_LOCK:
        return _job_snapshot(SCAN_JOBS[job_id])


@router.get("/scan/resume")
def resume_scan_job(job_id: str = Query(...)):
    with SCAN_JOBS_LOCK:
        job = SCAN_JOBS.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Scan job not found")
        if job["status"] in ("completed", "error", "cancelled"):
            return _job_snapshot(job)
        job["pause_requested"] = False
        job["status"] = "running"
        job["updated_at"] = datetime.utcnow().isoformat() + "Z"
    _job_log(job_id, "Scan resumed by user.")
    with SCAN_JOBS_LOCK:
        return _job_snapshot(SCAN_JOBS[job_id])


@router.get("/scan/cancel")
def cancel_scan_job(job_id: str = Query(...)):
    with SCAN_JOBS_LOCK:
        job = SCAN_JOBS.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Scan job not found")
        job["cancel_requested"] = True
        job["pause_requested"] = False
        job["status"] = "cancelled"
        job["error"] = "Scan cancelled"
        job["updated_at"] = datetime.utcnow().isoformat() + "Z"
    _job_log(job_id, "Scan cancelled by user.")
    with SCAN_JOBS_LOCK:
        return _job_snapshot(SCAN_JOBS[job_id])


@router.get("/download_report")
def download_report(path: str = Query(...), format: str = Query("json")):
    fmt = format.lower()
    if fmt not in ("json", "txt", "pdf", "docx"):
        raise HTTPException(status_code=400, detail="Unsupported format requested")

    report_path = _sanitize_report_path(path)

    # We need JSON source for PDF/DOCX conversion. For raw request, we can also return text or json directly.
    if fmt in ("json", "txt"):
        if not os.path.exists(report_path):
            raise HTTPException(status_code=404, detail="Report file not found")

        media_type = "application/json" if fmt == "json" else "text/plain"
        filename = _format_report_filename(report_path, fmt)

        return FileResponse(report_path, media_type=media_type, filename=filename)

    json_path = report_path
    if not json_path.lower().endswith(".json"):
        # prefer the json report if user passed txt
        json_path = report_path.rsplit(".", 1)[0] + ".json"

    json_path = _sanitize_report_path(json_path)

    if not os.path.exists(json_path):
        raise HTTPException(status_code=404, detail="Source JSON report not found")

    with open(json_path, "r", encoding="utf-8") as jf:
        payload = json.load(jf)

    out_filename = _format_report_filename(json_path, fmt)
    out_path = os.path.join("reports", f"tmp_{uuid.uuid4().hex}_{out_filename}")

    if fmt == "pdf":
        _generate_pdf_report(payload, out_path)
        media_type = "application/pdf"
    else:
        _generate_docx_report(payload, out_path)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return FileResponse(out_path, media_type=media_type, filename=out_filename)


@router.post("/surface/discover")
def discover_surface(payload: dict = Body(...)):
    target = str(payload.get("target") or "").strip()
    if not target:
        raise HTTPException(status_code=400, detail="target is required")
    return run_discover_web_surface(target, payload.get("auth") or {})


@router.post("/api-security/import")
def import_api_document(payload: dict = Body(...)):
    document = payload.get("document")
    if not document:
        raise HTTPException(status_code=400, detail="document is required")
    return run_analyze_api_document(
        document,
        payload.get("format"),
        payload.get("base_url"),
        payload.get("auth") or {},
    )


@router.post("/api-security/check")
def check_api_endpoints(payload: dict = Body(...)):
    endpoints = payload.get("endpoints") or []
    if not endpoints:
        raise HTTPException(status_code=400, detail="endpoints are required")
    return run_assess_api_endpoints(endpoints, payload.get("auth") or {})


@router.get("/scan/compare")
def compare_scan_reports(path_a: str = Query(...), path_b: str = Query(...)):
    safe_a = _sanitize_report_path(path_a)
    safe_b = _sanitize_report_path(path_b)
    if not os.path.exists(safe_a) or not os.path.exists(safe_b):
        raise HTTPException(status_code=404, detail="One or both report files were not found")
    return compare_reports(safe_a, safe_b)


@router.get("/schedules")
def get_scan_schedules():
    return {
        "schedules": list_schedules(),
        "runs": list_schedule_runs(),
    }


@router.post("/schedules")
def create_scan_schedule(payload: dict = Body(...)):
    target = str(payload.get("target") or "").strip()
    if not target:
        raise HTTPException(status_code=400, detail="target is required")

    schedule = save_schedule(
        target=target,
        interval_minutes=int(payload.get("interval_minutes") or 30),
        options=payload.get("options") or {},
    )
    schedule_scan_job(schedule, lambda job_target, options: _execute_scan(job_target, options=options))
    return schedule


@router.delete("/schedules/{schedule_id}")
def remove_scan_schedule(schedule_id: str):
    if not delete_schedule(schedule_id):
        raise HTTPException(status_code=404, detail="schedule not found")
    return {"deleted": True, "schedule_id": schedule_id}


@router.get("/stats")
def get_stats():
    """Aggregate statistics from all scan reports."""
    reports_dir = "reports"
    if not os.path.exists(reports_dir):
        return {
            "scan_history": [],
            "totals": {
                "scans": 0,
                "hosts_scanned": 0,
                "total_vulns": 0,
                "critical": 0,
                "exposed": 0,
                "ips": 0,
                "ranges": 0,
                "domains": 0,
                "avg_risk_score": 0
            },
            "common_ports": [],
            "vuln_breakdown": [],
            "os_stats": [],
            "risk_distribution": {
                "critical": 0,
                "high": 0,
                "medium": 0,
                "low": 0
            }
        }

    latest_scan_by_target = {}
    for filename in os.listdir(reports_dir):
        if not filename.endswith(".json") or filename.startswith("tmp_"):
            continue

        # Parse filename: scan_{target}_{timestamp}.json
        parts = filename.replace("scan_", "").replace(".json", "").rsplit("_", 2)
        if len(parts) != 3:
            continue

        target_str, date_str, time_str = parts
        timestamp_str = f"{date_str}_{time_str}"

        # Determine type
        if "/" in target_str or "-" in target_str:
            scan_type = "Range"
        elif is_domain_input(target_str):
            scan_type = "Domain"
        else:
            scan_type = "IP"

        filepath = os.path.join(reports_dir, filename)
        dt = None
        try:
            dt = datetime.fromtimestamp(os.path.getmtime(filepath))
            date_formatted = dt.strftime("%b %d, %I:%M %p").replace(f"{dt.strftime('%b')} 0", f"{dt.strftime('%b')} ")
        except Exception:
            try:
                dt = datetime.strptime(timestamp_str, "%Y%m%d_%H%M%S")
                date_formatted = dt.strftime("%b %d, %I:%M %p").replace(f"{dt.strftime('%b')} 0", f"{dt.strftime('%b')} ")
            except Exception:
                date_formatted = timestamp_str

        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
        except:
            continue

        original_target = str(data.get("input") or target_str).strip()
        if "/" in original_target or "-" in original_target:
            scan_type = "Range"
        elif is_domain_input(original_target):
            scan_type = "Domain"
        else:
            scan_type = "IP"

        assets = data.get("assets", [])
        vuln_summary = data.get("vulnerability_summary", {})
        active_hosts = data.get("active_hosts", len(assets))
        vulns = vuln_summary.get("total_vulnerabilities", 0)
        critical = vuln_summary.get("critical_risk", 0)

        # Calculate risk score (simplified)
        risk_score = min(100, (critical * 20) + (vulns * 5) + (active_hosts * 2))

        # Count open ports and retain the exact protocol/service labels seen in reports
        open_ports_count = 0
        port_counts = {}
        port_label_counts = {}
        os_counts = {}
        vuln_counts = {}
        for asset in assets:
            open_ports = asset.get("open_ports", [])
            open_ports_count += len(open_ports)
            for port_info in open_ports:
                port = port_info.get("port")
                if port:
                    port_counts[port] = port_counts.get(port, 0) + 1
                    service_name = str(port_info.get("service") or "").strip()
                    protocol_name = str(port_info.get("protocol") or "").strip()
                    if service_name and service_name.lower() != "unknown":
                        port_label = service_name.upper()
                    elif protocol_name:
                        port_label = protocol_name.upper()
                    else:
                        port_label = f"PORT {port}"
                    label_map = port_label_counts.setdefault(port, {})
                    label_map[port_label] = label_map.get(port_label, 0) + 1

            os_name = asset.get("os", "Unknown")
            if os_name and os_name != "Unknown":
                os_counts[os_name] = os_counts.get(os_name, 0) + 1

        # Build the scan record (including raw timestamp for sorting later)
        scan_record = {
            "id": 0,
            "target": original_target,
            "type": scan_type,
            "date": date_formatted,
            "hostsUp": active_hosts,
            "openPorts": open_ports_count,
            "vulns": vulns,
            "critical": critical,
            "status": "Complete",
            "riskScore": risk_score,
            "timestamp": dt if isinstance(dt, datetime) else None,
            "reportPath": filepath.replace("\\", "/")
        }

        # Count vulnerabilities by type (simplified - using service names)
        for asset in assets:
            for vuln in asset.get("vulnerabilities", []):
                vuln_type = vuln.get("title", "Other").split()[0]  # First word
                vuln_counts[vuln_type] = vuln_counts.get(vuln_type, 0) + 1

        target_key = original_target.strip().lower()
        existing = latest_scan_by_target.get(target_key)
        existing_ts = existing["record"].get("timestamp") if existing else None
        current_ts = scan_record.get("timestamp")

        if (
            existing is None
            or (current_ts or datetime.min) >= (existing_ts or datetime.min)
        ):
            latest_scan_by_target[target_key] = {
                "record": scan_record,
                "active_hosts": active_hosts,
                "vulns": vulns,
                "critical": critical,
                "risk_score": risk_score,
                "type": scan_type,
                "port_counts": port_counts,
                "port_label_counts": port_label_counts,
                "os_counts": os_counts,
                "vuln_counts": vuln_counts,
            }

    # Sort scan history by timestamp (newest first)
    latest_scan_entries = list(latest_scan_by_target.values())
    scan_history = [entry["record"] for entry in latest_scan_entries]
    scan_history.sort(key=lambda x: x.get("timestamp") or datetime.min, reverse=True)
    for idx, record in enumerate(scan_history, start=1):
        record["id"] = idx

    total_hosts = sum(entry["active_hosts"] for entry in latest_scan_entries)
    total_vulns = sum(entry["vulns"] for entry in latest_scan_entries)
    total_critical = sum(entry["critical"] for entry in latest_scan_entries)
    total_ips = sum(1 for entry in latest_scan_entries if entry["type"] == "IP")
    total_ranges = sum(1 for entry in latest_scan_entries if entry["type"] == "Range")
    total_domains = sum(1 for entry in latest_scan_entries if entry["type"] == "Domain")
    risk_scores = [entry["risk_score"] for entry in latest_scan_entries]

    # Calculate totals
    total_scans = len(scan_history)
    exposed = sum(1 for s in scan_history if s["riskScore"] >= 70)
    avg_risk_score = int(sum(risk_scores) / len(risk_scores)) if risk_scores else 0

    totals = {
        "scans": total_scans,
        "hosts_scanned": total_hosts,
        "total_vulns": total_vulns,
        "critical": total_critical,
        "exposed": exposed,
        "ips": total_ips,
        "ranges": total_ranges,
        "domains": total_domains,
        "avg_risk_score": avg_risk_score
    }

    all_ports = {}
    all_port_labels = {}
    all_os = {}
    all_vulns = {}
    for entry in latest_scan_entries:
        for port, count in entry["port_counts"].items():
            all_ports[port] = all_ports.get(port, 0) + count
        for port, labels in entry["port_label_counts"].items():
            label_map = all_port_labels.setdefault(port, {})
            for label, count in labels.items():
                label_map[label] = label_map.get(label, 0) + count
        for os_name, count in entry["os_counts"].items():
            all_os[os_name] = all_os.get(os_name, 0) + count
        for vuln_name, count in entry["vuln_counts"].items():
            all_vulns[vuln_name] = all_vulns.get(vuln_name, 0) + count

    # Common ports - top 10
    sorted_ports = sorted(all_ports.items(), key=lambda x: x[1], reverse=True)[:10]
    total_port_hosts = sum(all_ports.values())
    common_ports = []
    for port, count in sorted_ports:
        pct = int((count / total_port_hosts) * 100) if total_port_hosts > 0 else 0
        labels = all_port_labels.get(port, {})
        service = max(labels.items(), key=lambda item: item[1])[0] if labels else f"PORT {port}"
        color = ["#059669", "#1d4ed8", "#7c3aed", "#d97706", "#0891b2", "#dc2626", "#6d28d9", "#be185d", "#4338ca", "#0d9488"][len(common_ports) % 10]
        common_ports.append({
            "port": port,
            "service": service,
            "count": count,
            "pct": pct,
            "color": color
        })

    # Vulnerability breakdown
    sorted_vulns = sorted(all_vulns.items(), key=lambda x: x[1], reverse=True)[:5]
    total_vuln_count = sum(all_vulns.values())
    vuln_breakdown = []
    colors = ["#ef4444", "#f97316", "#eab308", "#84cc16", "#6b7280"]
    for i, (name, count) in enumerate(sorted_vulns):
        pct = int((count / total_vuln_count) * 100) if total_vuln_count > 0 else 0
        vuln_breakdown.append({
            "name": name,
            "count": count,
            "pct": pct,
            "color": colors[i % len(colors)]
        })

    # OS stats
    sorted_os = sorted(all_os.items(), key=lambda x: x[1], reverse=True)
    total_os_hosts = sum(all_os.values())
    os_colors = ['#00e5ff', '#818cf8', '#475569']  # Cyan, Blue, Gray
    os_stats = []
    for i, (name, hosts) in enumerate(sorted_os[:3]):
        pct = int((hosts / total_os_hosts) * 100) if total_os_hosts > 0 else 0
        os_stats.append({
            "name": name,
            "pct": pct,
            "hosts": hosts,
            "color": os_colors[i]
        })
    # Add Unknown if not present
    if not any(os["name"] == "Unknown" for os in os_stats) and total_os_hosts < total_hosts:
        unknown_hosts = total_hosts - total_os_hosts
        pct = int((unknown_hosts / total_hosts) * 100) if total_hosts > 0 else 0
        os_stats.append({
            "name": "Unknown",
            "pct": pct,
            "hosts": unknown_hosts,
            "color": "#6b7280"
        })

    # Risk distribution
    risk_distribution = {
        "critical": sum(1 for s in scan_history if s["riskScore"] >= 70),
        "high": sum(1 for s in scan_history if 40 <= s["riskScore"] < 70),
        "medium": sum(1 for s in scan_history if 20 <= s["riskScore"] < 40),
        "low": sum(1 for s in scan_history if s["riskScore"] < 20)
    }

    return {
        "scan_history": scan_history,
        "totals": totals,
        "common_ports": common_ports,
        "vuln_breakdown": vuln_breakdown,
        "os_stats": os_stats,
        "risk_distribution": risk_distribution
    }
